import logging
from pathlib import Path
from typing import Any, Dict, Iterable, List

from docx import Document
from docx.shared import Pt

logger = logging.getLogger(__name__)


def _safe_list(value: Any) -> List[Any]:
    return value if isinstance(value, list) else []


def _add_heading(doc: Document, title: str) -> None:
    doc.add_heading(title, level=1)


def _add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        text = str(item).strip()
        if text:
            doc.add_paragraph(text, style="List Bullet")


def generate_resume_docx(resume_payload: Dict[str, Any], output_path: str) -> str:
    """Generate ATS-friendly resume DOCX with simple formatting."""
    path = Path(output_path)
    path.parent.mkdir(parents=True, exist_ok=True)

    logger.info("Generating DOCX at %s", path)
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    contact_info = resume_payload.get("contact_info", {})
    name = contact_info.get("name", "Candidate")
    doc.add_heading(str(name), level=0)

    contact_bits = [
        contact_info.get("email", ""),
        contact_info.get("phone", ""),
        contact_info.get("location", ""),
        contact_info.get("linkedin", ""),
        contact_info.get("portfolio", ""),
    ]
    contact_line = " | ".join([bit.strip() for bit in contact_bits if str(bit).strip()])
    if contact_line:
        doc.add_paragraph(contact_line)

    _add_heading(doc, "Professional Summary")
    doc.add_paragraph(str(resume_payload.get("professional_summary", "")).strip())

    _add_heading(doc, "Skills")
    _add_bullets(doc, _safe_list(resume_payload.get("skills")))

    _add_heading(doc, "Experience")
    for role in _safe_list(resume_payload.get("experience")):
        title = role.get("title", "")
        company = role.get("company", "")
        dates = role.get("dates", "")
        header = " - ".join([part for part in [title, company] if str(part).strip()])
        if dates:
            header = f"{header} ({dates})" if header else str(dates)
        if header:
            doc.add_paragraph(header)
        _add_bullets(doc, _safe_list(role.get("bullets")))

    projects = _safe_list(resume_payload.get("projects"))
    if projects:
        _add_heading(doc, "Projects")
        for project in projects:
            name_line = str(project.get("name", "")).strip()
            tech = ", ".join(_safe_list(project.get("tech")))
            if tech and name_line:
                doc.add_paragraph(f"{name_line} | {tech}")
            elif name_line:
                doc.add_paragraph(name_line)
            _add_bullets(doc, _safe_list(project.get("bullets")))

    _add_heading(doc, "Education")
    for edu in _safe_list(resume_payload.get("education")):
        degree = edu.get("degree", "")
        institution = edu.get("institution", "")
        dates = edu.get("dates", "")
        line = " - ".join([part for part in [degree, institution] if str(part).strip()])
        if dates:
            line = f"{line} ({dates})" if line else str(dates)
        if line:
            doc.add_paragraph(line)
        _add_bullets(doc, _safe_list(edu.get("details")))

    doc.save(path)
    return str(path.resolve())
